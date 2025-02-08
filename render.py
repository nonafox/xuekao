import docx
from docx import document
from docx.enum.text import *
from docx.enum.section import *
from docx.enum.table import *
from docx.shared import *
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import conf
import util
import random

def cell_center(cell, bold = False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    para = cell.paragraphs[0]
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = para.runs[0]
    run.font.size = Pt(14) if bold else Pt(12)
    run.font.bold = bold
def cell_format(cell, bold = False):
    cell_center(cell, bold)

def render_1(data: list):
    doc = docx.Document()
    assert isinstance(doc, document.Document)
    sec = doc.sections[0]
    sec.orientation = WD_ORIENTATION.LANDSCAPE
    sec.top_margin = conf.page_margin_top
    sec.left_margin = conf.page_margin_left
    sec.bottom_margin = conf.page_margin_bottom
    sec.right_margin = conf.page_margin_right
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    doc.styles['Normal'].font.name = '宋体'
    
    virgin = 1
    j = 0
    gen_j_dig = len(str(len(data)))
    for room in data:
        if virgin:
            virgin = 0
        else:
            doc.add_page_break()
        title = doc.add_paragraph('考场 %s      人数 %d' % (str(j + 1).rjust(gen_j_dig, '0'), len(room)))
        title.runs[0].font.size = Pt(15)
        title.runs[0].font.bold = 1
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        table = doc.add_table(conf.gen_rows, conf.gen_cols)
        table.style = 'Table Grid'

        i = 1
        orient = 1
        row, col = 0, conf.gen_cols - 1
        for student in room:
            cell = table.cell(row, col)
            cell.text = str(i).rjust(conf.gen_i_dig, '0')
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            p = cell.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            l = p.add_run('%s %s号\n' % (student[conf.key_class], util.str_select(str(student[conf.key_sid]), '？')))
            l.font.size = Pt(9)
            l = p.add_run('%s\n' % student[conf.key_name])
            l.font.size = Pt(12)
            l.font.bold = 1
            l = p.add_run('考号 %s\n' % util.str_select(student[conf.key_eid], '未知'))
            l.font.size = Pt(9)
            row += orient
            if row < 0:
                row = 0
                col -= 1
                orient *= - 1
            elif row > conf.gen_rows - 1:
                row = conf.gen_rows - 1
                col -= 1
                orient *= - 1
            i += 1
        j += 1
    doc.save(conf.path_out_1)
def render_2(data: dict):
    doc = docx.Document()
    assert isinstance(doc, document.Document)
    sec = doc.sections[0]
    sec.top_margin = conf.page_margin_top
    sec.left_margin = conf.page_margin_left
    sec.bottom_margin = conf.page_margin_bottom
    sec.right_margin = conf.page_margin_right
    doc.styles['Normal'].font.name = '宋体'
    virgin = 1

    for cl_name, cl in data.items():
        if virgin:
            virgin = 0
        else:
            doc.add_page_break()
        title = doc.add_paragraph(cl_name)
        title.runs[0].font.size = Pt(15)
        title.runs[0].font.bold = 1
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        table = doc.add_table(0, len(conf.columns))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        header = table.add_row()
        for j, col in enumerate(conf.columns):
            cell = header.cells[j]
            cell.text = col
            cell_format(cell, True)
        
        for student in cl:
            row = table.add_row()
            for j, col in enumerate(conf.columns):
                cell = row.cells[j]
                cell.text = student[col]
                cell_format(cell)
    doc.save(conf.path_out_2)

def render(data: list):
    data_splited = data.copy()
    random.shuffle(data_splited)
    data_splited = util.split_array(data_splited, conf.gen_rows * conf.gen_cols)
    for i, cl in enumerate(data_splited):
        for student in cl:
            student[conf.key_new_class] = str(i + 1)
    render_1(data_splited)

    data_classified = {}
    for student in data:
        if student[conf.key_class] not in data_classified:
            data_classified[student[conf.key_class]] = []
        cl = data_classified[student[conf.key_class]]
        cl.append(student)
    render_2(data_classified)
